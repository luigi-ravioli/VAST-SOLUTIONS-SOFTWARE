import pandas as pd
import tkinter
from tkinter import messagebox
from tkinter import *
from docx import Document

# Days Worked Options
days_options = [
    0,
    1,
    2,
    3,
    4,
    5,
    6,
    7
]

employees_list = []
employees_rates = []
processed_employees_list = []

# Initialize List of Employees
def InitializeEmployees():
    # Get list of employees from text file
    employees_file = open("PAYROLL\employees.txt", "r")
    for employee in employees_file:
        employeeListEntry = []
        employeeRateEntry = []
        rateDict = {}
        readLine = employee.strip().split(" ")
        employeeListEntry.append(readLine[0])
        employeeListEntry.append(readLine[1])
        employeeListEntry.append(readLine[2])
        rateDict["Daily Rate"] = float(readLine[3])
        rateDict["Overtime Rate"] = float(readLine[4])
        employeeRateEntry.append(readLine[0])
        employeeRateEntry.append(rateDict)
        employees_list.append(employeeListEntry)
        employees_rates.append(employeeRateEntry)
        
    employees_file.close()

    # Sort employee lists according to employee number
    employees_list.sort(key=lambda x: x[0])
    employees_rates.sort(key=lambda x: x[0])
    #dailyRateLabel.config(text=f"Daily Rate: {employees_rates[1]['Daily Rate']}")
    #overtimeRateLabel.config(text=f"Overtime Rate: {employees_rates[1]['Overtime Rate']}")

def UpdateDisplayedRates(*args):
    employeeNumber = employeeVar.get().replace("(","").replace("'", "").replace(")", "").split(",")[0].strip()
    print("Attempting to update displayed rates!")
    for employee in employees_rates:
        if employeeNumber in employee:
            dailyRateLabel.config(text=f"Daily Rate: {employee[1]['Daily Rate']}")
            overtimeRateLabel.config(text=f"Overtime Rate: {employee[1]['Overtime Rate']}")
            print("Updated displayed rates!")

# Process Employee with Input Data
def ProcessEmployee():
    if employees_list: # Check if there are employees left to process
        #if len(dailyRateEntry.get()) != 0: # Check if user entered daily rate
            #if len(overtimeRateEntry.get()) != 0: # Check if user entered overtime rate
                # Place user input into an employee entry
                employee = []
                employeeNumber = employeeVar.get().replace("(","").replace("'", "").replace(")", "").split(",")[0].strip()
                lastName = employeeVar.get().replace("(","").replace("'", "").replace(")", "").split(",")[1].strip()
                firstName = employeeVar.get().replace("(","").replace("'", "").replace(")", "").split(",")[2].strip()
                employee.append(lastName) # Last Name
                employee.append(firstName) # First Name
                for row in employees_rates:
                    if employeeNumber in row:
                        employee.append(row[1]['Daily Rate']) # Daily Rate
                        employee.append(row[1]['Overtime Rate']) # Overtime Rate
                employee.append(daysWorkedVar.get()) # Days Worked
                employee.append(overtimeHoursEntry.get()) # Overtime Hours Rendered
                


                # Process employee entry
                processed_employees_list.append(employee) # Add entry to list of processed employees
                # Remove processed employee from employee list
                for row in employees_list:
                    if lastName in row and firstName in row:
                        print(f'Employee {lastName} {firstName} deleted!')
                        del employees_list[employees_list.index(row)]
                        break
                for row in employees_rates:
                    if employeeNumber in row:
                        print(f'Employee {employeeNumber} deleted!')
                        del employees_rates[employees_rates.index(row)]
                        break

                # Clear user input
                #dailyRateEntry.delete(0,END)
                #overtimeRateEntry.delete(0,END)
                daysWorkedVar.set(0)
                overtimeHoursEntry.delete(0,END)
                overtimeHoursEntry.insert(0, 0)

                # Update displayed processed employees and dropdown menu
                UpdateEmployees()
                UpdateOptions()

                # Check if there are employees left
                if employees_list:
                    # Update displayed employee to process to next employee
                    employeeVar.set(employees_list[0])

                # If all employees have been processed
                else:
                    # Update displayed employee to N/A
                    employeeVar.set("N/A")

            #else:
                #messagebox.showinfo("Error", "Please enter a value for Overtime Rate")
        #else:
            #messagebox.showinfo("Error", "Please enter a value for Daily Rate")
    #else:
        #messagebox.showinfo("Error", "All employees have been processed")

# Update Options in Dropdown Menu
def UpdateOptions():
    employeeEntry['menu'].delete(0, 'end') # Delete all options

    # Iterate through all employees
    for employee in employees_list:
        employeeEntry['menu'].add_command(label=employee, command=tkinter._setit(employeeVar, employee)) # Repopulate options in dropdown menu

    # Update displayed rates
    #dailyRateLabel.config(text=f"Daily Rate: {employees_rates[0][1]['Daily Rate']}")
    #overtimeRateLabel.config(text=f"Overtime Rate: {employees_rates[0][1]['Overtime Rate']}")
    print("Refreshed options!")

# Update Table of Processed Employees
def UpdateEmployees():
    currentRow = len(processed_employees_list) # Get most recent index from list of processed employees
    # Iterate through employee's data
    for currentCol in range(6): 
        # Generate labels to display employee's data
        employeeData = Label(resultFrame)
        employeeData.config(text=processed_employees_list[currentRow-1][currentCol])
        employeeData.grid(row=currentRow, column=currentCol)
        print(processed_employees_list[currentRow-1][currentCol])

# Calculate Employee Salary
def CalculateSalary(daily_rate, days_worked, overtime_rate, overtime_hours):
    # Calculate base, overtime and total salary
    base_salary = daily_rate * days_worked
    overtime_salary = overtime_rate * overtime_hours
    total_salary = base_salary + overtime_salary

    return base_salary, overtime_salary, total_salary

# Compute Payslips of Employees
def ComputePayslip():
    payslip_data = pd.DataFrame() # Dataframe to store payslip data
    document = Document() # Word document

    if processed_employees_list: # Check if there are processed employees
        # Iterate through each processed employee
        for employee in processed_employees_list:
            # Get employee data
            lastName = employee[0]
            firstName = employee[1]
            dailyRate = float(employee[2])
            overtimeRate = float(employee[3])
            daysWorked = float(employee[4])
            overtimeHours = float(employee[5])

            # Calculate salaries
            baseSalary, overtimeSalary, totalSalary = CalculateSalary(dailyRate, daysWorked, overtimeRate, overtimeHours)

            # Create a dictionary and fill in with payslip data
            payslip = {
                'Last Name': lastName,
                'First Name': firstName,
                'Daily Rate': dailyRate,
                'Overtime Rate': overtimeRate,
                'Days Worked': daysWorked,
                'Overtime Hours': overtimeHours,
                'Base Salary': baseSalary,
                'Overtime Salary': overtimeSalary,
                'Total Salary': totalSalary
            }

            # Create a new page for each employee and fill with employee data
            document.add_paragraph(f"Last Name: {lastName}")
            document.add_paragraph(f"First Name: {firstName}")
            document.add_paragraph(f"Daily Rate: {dailyRate}")
            document.add_paragraph(f"Overtime Rate: {overtimeRate}")
            document.add_paragraph(f"Days Worked: {daysWorked}")
            document.add_paragraph(f"Overtime Hours: {overtimeHours}")
            document.add_paragraph(f"Base Salary: {baseSalary}")
            document.add_paragraph(f"Overtime Salary: {overtimeSalary}")
            document.add_paragraph(f"Total Salary: {totalSalary}")
            document.add_page_break()

            # Append payslip data into Dataframe
            payslip_data = payslip_data.append(payslip, ignore_index=True)

        # Generate Excel file and place payslip data
        excel_output_file = 'payslips.xlsx'
        word_output_file = 'payslips.docx'
        payslip_data.to_excel(excel_output_file, index=False)
        document.save(word_output_file)
        messagebox.showinfo("Success", f"Payslips generated and saved in {excel_output_file} and {word_output_file}.")
    
    else:
        messagebox.showinfo("Error", "No employees have been processed.")

# Initialize
InitializeEmployees()

# Main GUI Window
window = Tk()
window.title("Vast Solutions Payroll Calculator")
window.configure(background="white smoke", padx=10)
window.geometry("700x700")

# Payroll Result Frame
resultFrame = Frame(window, bg="white smoke", padx=10, pady=10)
resultFrame.grid(row=1, column=0, sticky=W)
lastNameTitle = Label(resultFrame, text="Last Name", bg="white smoke", fg="black", font="none 12 bold")
lastNameTitle.grid(row=0, column=0, sticky=W)
firstNameTitle = Label(resultFrame, text="First Name", bg="white smoke", fg="black", font="none 12 bold")
firstNameTitle.grid(row=0, column=1, sticky=W)
dailyRateTitle = Label(resultFrame, text="Daily Rate", bg="white smoke", fg="black", font="none 12 bold")
dailyRateTitle.grid(row=0, column=2, sticky=W)
overtimeRateTitle = Label(resultFrame, text="Overtime Rate", bg="white smoke", fg="black", font="none 12 bold")
overtimeRateTitle.grid(row=0, column=3, sticky=W)
daysWorkedTitle = Label(resultFrame, text="Days Worked", bg="white smoke", fg="black", font="none 12 bold")
daysWorkedTitle.grid(row=0, column=4, sticky=W)
overtimeHoursTitle = Label(resultFrame, text="Overtime Hours", bg="white smoke", fg="black", font="none 12 bold")
overtimeHoursTitle.grid(row=0, column=5, sticky=W)

# Compute Payslip Button
computePayslipButton = Button(window, text="Generate Payslip", command=ComputePayslip, relief="raised")
computePayslipButton.grid(row=2, column=0, sticky=W)

# Employee Processing Frame
employeeProcessingFrame = Frame(window, bg="white smoke", pady=10)
employeeProcessingFrame.grid(row=0, column=0, sticky=W)

# Title Frame
titleFrame = Frame(employeeProcessingFrame, bg="white smoke")
titleFrame.grid(row=0, column=0, sticky=W)
titleLabel = Label(titleFrame, text="Payroll Calculator", bg="white smoke", fg="black", font="none 12 bold")
titleLabel.grid(row=0, column=0, sticky=W)

# Employee Frame
employeeFrame = Frame(employeeProcessingFrame, bg="white smoke")
employeeFrame.grid(row=1, column=0, sticky=W)
employeeLabel = Label(employeeFrame, text="Employee: ", bg="white smoke", fg="black", font="none 10")
employeeLabel.grid(row=0, column=0, sticky=W)
employeeVar = StringVar()
employeeVar.trace("w", UpdateDisplayedRates)
employeeEntry = OptionMenu(employeeFrame, employeeVar, *employees_list)
employeeEntry.grid(row=0, column=1, sticky=W)

# Daily Rate Input Frame
dailyRateFrame = Frame(employeeProcessingFrame, bg="white smoke")
dailyRateFrame.grid(row=2, column=0, sticky=W)
dailyRateLabel = Label(dailyRateFrame, text="Daily Rate:", bg="white smoke", fg="black", font="none 10")
dailyRateLabel.grid(row=0, column=0, sticky=W)

# Overtime Hourly Rate Input Frame
overtimeRateFrame = Frame(employeeProcessingFrame, bg="white smoke")
overtimeRateFrame.grid(row=3, column=0, sticky=W)
overtimeRateLabel = Label(overtimeRateFrame, text="Overtime Hourly Rate:", bg="white smoke", fg="black", font="none 10")
overtimeRateLabel.grid(row=0, column=0, sticky=W)

# Days Worked Input Frame
daysWorkedFrame = Frame(employeeProcessingFrame, bg="white smoke")
daysWorkedFrame.grid(row=4, column=0, sticky=W)
daysWorkedLabel = Label(daysWorkedFrame, text="Enter Days Worked this Week:", bg="white smoke", fg="black", font="none 10")
daysWorkedLabel.grid(row=0, column=0, sticky=W)
daysWorkedVar = IntVar()
daysWorkedVar.set(0)
daysWorkedEntry = OptionMenu(daysWorkedFrame, daysWorkedVar, *days_options)
daysWorkedEntry.grid(row=0, column=1, sticky=W)

# Overtime Hours Input Frame
overtimeHoursFrame = Frame(employeeProcessingFrame, bg="white smoke")
overtimeHoursFrame.grid(row=5, column=0, sticky=W)
overtimeHoursLabel = Label(overtimeHoursFrame, text="Enter Overtime Hours:", bg="white smoke", fg="black", font="none 10")
overtimeHoursLabel.grid(row=0, column=0, sticky=W)
overtimeHoursEntry = Entry(overtimeHoursFrame, width=10, bg="white smoke", fg="black", font="none 8")
overtimeHoursEntry.insert(0, 0)
overtimeHoursEntry.grid(row=0, column=1, sticky=W)

# Submit Button
submitButton = Button(employeeProcessingFrame, text="Submit", command=ProcessEmployee, relief="raised")
submitButton.grid(row=6, column=0, sticky=W)

window.mainloop()
